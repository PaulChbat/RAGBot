# Functions for file_page
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
import os
import streamlit as st
from langchain_community.document_loaders import PyPDFLoader
import pandas as pd
from docx import Document as DocxDocument

# Class used to format text loaded from files to be compatible with the chunking function of langchain
class Document:
    def __init__(self, page_content, metadata):
        self.page_content = page_content
        self.metadata = metadata

def init_folder (folder):
    os.makedirs(folder, exist_ok=True)
    # Clear files in the folder
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        if os.path.isfile(file_path):
            os.remove(file_path)

def save_file(file_path, uploaded_file):
    with open(file_path, "wb") as f:
        f.write(uploaded_file.read())

def setup_vdb(used_model_name,texts):
    embeddings = HuggingFaceEmbeddings(model_name=used_model_name)
    vectorstore = FAISS.from_documents(texts, embeddings)
    vectorstore.save_local('db_faiss')
    st.session_state['vectorstore'] = vectorstore

def show_content(documents):
    st.subheader("Content")
    with st.expander("Click to expand"):
        for doc in documents:
            st.text(doc.page_content)

def show_chunks(texts):
    st.subheader("Chunks")
    chunk_index = 1
    with st.expander("Click to expand"):
        for text in texts:
            st.write(chunk_index)
            st.text(text)
            chunk_index += 1

# Function to save uploaded files locally and load them in documents
def save_load_files(uploaded_files):
    folder = "Files"
    # Either create the folder 
    init_folder(folder)
    documents = []

    for uploaded_file in uploaded_files:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        file_path = os.path.join(folder, uploaded_file.name)
        save_file(file_path, uploaded_file)

        if file_extension == "pdf":
            loader = PyPDFLoader(file_path)
            pdf_content = loader.load()
            for i, doc in enumerate(pdf_content):
                documents.append(Document(doc.page_content, {"source": uploaded_file.name, "page": i + 1}))

        elif file_extension == "docx":
            doc = DocxDocument(file_path)
            doc_content = "\n".join([para.text for para in doc.paragraphs])
            documents.append(Document(doc_content, {"source": uploaded_file.name}))

        elif file_extension == "xlsx":
            df = pd.read_excel(file_path)
            xl_content = df.to_string(index=False)
            documents.append(Document(xl_content, {"source": uploaded_file.name}))

        else:
            raise ValueError(f"Unsupported file type:  {file_extension}")
    
    return documents

def show_files_in_database(vectorstore):
    st.subheader("Files in the DataBase")
    with st.expander("Click to expand"):
        # Collect all unique source filenames
        sources = set()
        for doc in vectorstore.docstore._dict.values():
            metadata = doc.metadata
            source = metadata.get("source", "Unknown")
            sources.add(source)  # Add to the set to avoid duplicates
        # Display the unique filenames
        for source in sources:
            st.write(source)

# Functions for chatbot_page
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from pydantic import BaseModel, Field
from langchain.llms.base import LLM
from typing import Any, List, Optional, Dict
import os
from groq import Groq
from langchain_community.callbacks.manager import get_openai_callback
from langchain.chains import RetrievalQA
import pyaudio
import wave
from deepgram import (DeepgramClient, SpeakOptions)


class GroqLLMConfig(BaseModel):
    model_name: str = Field(..., description="The name of the Groq model to use.")
    temperature: float = Field(0.0, description="The temperature to use for sampling.")
    groq_api_key: str = Field(..., description="The API key for Groq.")

class GroqLLM(LLM):
    config: GroqLLMConfig
    client: Any = None

    def __init__(self, model_name: str, temperature: float = 0.0, groq_api_key: Optional[str] = None):
        super().__init__()
        groq_api_key = groq_api_key
        if not groq_api_key:
            raise ValueError("Groq API key must be provided or set as GROQ_API_KEY environment variable.")

        self.config = GroqLLMConfig(
            model_name=model_name,
            temperature=temperature,
            groq_api_key=groq_api_key
        )
        self.client = Groq(api_key=self.config.groq_api_key)

    @property
    def config(self) -> GroqLLMConfig:
        return self._config

    @config.setter
    def config(self, value: GroqLLMConfig):
        self._config = value

    def _call(self, prompt: str, stop: Optional[List[str]] = None) -> str:
        response = self.client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model=self.config.model_name,
            temperature=self.config.temperature,
        )
        return response.choices[0].message.content

    @property
    def _llm_type(self) -> str:
        return "Groq"

    @property
    def _identifying_params(self) -> Dict[str, Any]:
        return {"model_name": self.config.model_name, "temperature": self.config.temperature}
    
# Create a Groq LLM instance
llm = GroqLLM(
    model_name="llama3-8b-8192",  # "llama-3.1-70b-versatile"
    temperature=0.3,
    groq_api_key=os.getenv("GROQ_API_KEY")
)
# Define a custom prompt template
template = """Use the following pieces of context to answer the question at the end. If you don't know the answer, just say that you don't know, don't try to make up an answer.

{context}

Question: {question}
Answer: """

PROMPT = PromptTemplate(
    template=template, input_variables=["context", "question"]
)

def setup_qa(vectorstore):
    qa = RetrievalQA.from_chain_type(
    llm=llm,
    chain_type="stuff",
    retriever=vectorstore.as_retriever(),
    return_source_documents=True,
    chain_type_kwargs={"prompt": PROMPT}
    )
    return qa

def get_answer(vectorstore):
    # Retrieve the current chat history
    chat_history = st.session_state[st.session_state['current_chat']]

    # Format the chat history as a string to pass to the model
    history_str = ""
    for message in chat_history:
        if message["role"] == "user":
            history_str += f"User: {message['content']}\n"
        else:
            history_str += f"Assistant: {message['content']}\n"

    # Create the prompt by combining history with the current query
    full_prompt = f"{history_str}Assistant:"

    # Set up the QA chain
    qa = setup_qa(vectorstore)
    
    # Fetch the result with the entire conversation context
    with get_openai_callback() :
        result = qa({"query": full_prompt})  # Using the history as part of the query

    answer = result['result']
    source_documents = result['source_documents']
    
    response = (f"{answer}\n\n")

    # Initialize a set to keep track of unique source information
    unique_sources = set()

    for doc in source_documents:
        source_info = f"- {doc.metadata['source']}"
        if 'page' in doc.metadata:  # Only pdfs have pages.
            source_info += f", Page {doc.metadata['page']}."

        # Add the source info to the set if it's not already present
        if source_info not in unique_sources:
            unique_sources.add(source_info)

    # Now, add the unique sources to the response string
    if "I don't know" not in answer: 
        response += "Sources:\n"
        for source in unique_sources:
            response += f"{source}\n"

    return response

def text_to_speech(text, audio_path):
    sourceless_text = text.split("Sources")[0].strip() #remove sources from TTS 
    deepgram = DeepgramClient(os.environ.get("DEEPGRAM_API_KEY"))
    options = SpeakOptions(model='aura-asteria-en')
    deepgram.speak.v("1").save(audio_path, {"text":sourceless_text}, options)

def transcribe_audio(audio_file):
    client = Groq()
    with open(audio_file, "rb") as file:
        # Create a transcription of the audio file
        transcription = client.audio.transcriptions.create(
          file=(audio_file, file.read()), # Required audio file
          model="whisper-large-v3-turbo", # Required model to use for transcription
          prompt="Specify context or spelling",  # Change it to adapt to the context of the app
          response_format="json",  
          #language="en",  # Omitt to let Groq guess the language or you can specify one. 
          temperature=0.0  # The lower the temperature the more accurate the transcription is
        )
        return transcription.text

# Audio recording settings
FORMAT = pyaudio.paInt16  # Sets the audio format to a 16-bit signed integer
CHANNELS = 1  # Sets to mono channel recording, better for voice recognition
RATE = 44100  # Sample rate in Hz
CHUNK = 1024  # Chunk size, audio read in chunks of 1024 at a time
OUTPUT_FILE = "Audio/question.wav"  # Output file

# Initialize PyAudio
p = pyaudio.PyAudio()

# Start recording function
def start_recording():
    st.session_state['recording'] = True
    st.session_state['frames'] = []

    stream = p.open(format=FORMAT, channels=CHANNELS,
                    rate=RATE, input=True,
                    frames_per_buffer=CHUNK)

    #with st.spinner("Recording audio..."):
    while st.session_state['recording']:
        data = stream.read(CHUNK)
        st.session_state['frames'].append(data)

    # Stop and close the stream
    stream.stop_stream()
    stream.close()
    

# Stop recording function
def stop_recording():
    st.session_state['recording'] = False
    
    # Save the recorded data as a WAV file
    wf = wave.open(OUTPUT_FILE, 'wb')
    wf.setnchannels(CHANNELS)
    wf.setsampwidth(p.get_sample_size(FORMAT))
    wf.setframerate(RATE)
    wf.writeframes(b''.join(st.session_state['frames']))
    wf.close()

def get_audio_query():
    if os.path.exists("Audio/question.wav"):
        audio_query = transcribe_audio("Audio/question.wav")
        os.remove("Audio/question.wav") # Cleanup the audio file because i will have no use.
        return audio_query
    else:
        return None




    



