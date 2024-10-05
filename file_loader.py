import streamlit as st
import os
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from docx import Document as DocxDocument
from langchain_huggingface import HuggingFaceEmbeddings
import pandas as pd

# Class used to format text loaded from files to be compatible with the chunking function of langchain
class Document:
    def __init__(self, page_content, metadata):
        self.page_content = page_content
        self.metadata = metadata

def init_folder (folder):
    os.makedirs(folder, exist_ok=True)
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        if os.path.isfile(file_path):
            os.remove(file_path)

def save_file(file_path, uploaded_file):
    with open(file_path, "wb") as f:
        f.write(uploaded_file.read())

def setup_vdb(used_model_name,texts):
    embeddings = HuggingFaceEmbeddings(model_name=used_model_name)
    vectorstore = FAISS.from_documents(texts, embeddings)
    vectorstore.save_local('db_faiss')
    st.session_state['vectorstore'] = vectorstore

# Function to save uploaded files locally and load them in documents
def save_load_files(uploaded_files):
    folder = "Files"
    init_folder(folder)
    
    documents = []

    for uploaded_file in uploaded_files:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        file_path = os.path.join(folder, uploaded_file.name)
        save_file(file_path, uploaded_file)

        if file_extension == "pdf":
            loader = PyPDFLoader(file_path)
            pdf_content = loader.load()
            for i, doc in enumerate(pdf_content):
                documents.append(Document(doc.page_content, {"source": uploaded_file.name, "page": i + 1}))

        elif file_extension == "docx":
            doc = DocxDocument(file_path)
            doc_content = "\n".join([para.text for para in doc.paragraphs])
            documents.append(Document(doc_content, {"source": uploaded_file.name}))

        elif file_extension == "xlsx":
            df = pd.read_excel(file_path)
            xl_content = df.to_string(index=False)
            documents.append(Document(xl_content, {"source": uploaded_file.name}))

        else:
            raise ValueError(f"Unsupported file type:  {file_extension}")
    
    return documents

def file_page():
    st.title("Load your File")

    uploaded_files = st.file_uploader("", type=["pdf", "docx", "xlsx"], accept_multiple_files=True)
    
    if uploaded_files:
        for uploaded_file in uploaded_files:
            st.success(f"File uploaded: {uploaded_file.name}")
        
        documents = save_load_files(uploaded_files)

        st.subheader("Content")
        with st.expander("Click to expand"):
            for doc in documents:
                st.text(doc.page_content)

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=0)
        texts = text_splitter.split_documents(documents)

        st.subheader("Chunks")
        chunk_index = 1
        with st.expander("Click to expand"):
            for text in texts:
                st.write(chunk_index)
                st.text(text)
                chunk_index += 1

        used_model_name = "sentence-transformers/paraphrase-multilingual-mpnet-base-v2"
        setup_vdb(used_model_name,texts)
    st.write("---")
    if st.button("Load Last DataBase"):
        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/paraphrase-multilingual-mpnet-base-v2")
        try:
            vectorstore = FAISS.load_local('db_faiss', embeddings, allow_dangerous_deserialization=True)
            st.session_state['vectorstore'] = vectorstore
            st.success("Vectorstore loaded successfully!")

            # Show metadata (source filenames and page numbers) for the documents in the vectorstore
            if 'vectorstore' in st.session_state:
                st.subheader("Files Stored in Vectorstore")
                with st.expander("Click to expand"):
                    # Collect all unique source filenames
                    sources = set()
                    for doc in vectorstore.docstore._dict.values():
                        metadata = doc.metadata
                        source = metadata.get("source", "Unknown")
                        sources.add(source)  # Add to the set to avoid duplicates

                    # Display the unique filenames
                    for source in sources:
                        st.write(source)


        except FileNotFoundError:
            st.error("No saved vectorstore found! Please upload your files first.")
